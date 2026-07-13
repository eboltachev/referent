from docx import Document
from src.application.services.docx_service import format_mmss
def build_docx(path, job, aliases, segments):
    doc=Document(); doc.add_heading('Интеллектуальная Стенограмма',0)
    doc.add_paragraph(f'Название файла: {job.display_filename}')
    doc.add_paragraph(f'Дата загрузки/записи: {job.uploaded_at}')
    doc.add_paragraph(f'Найденные языки: {", ".join(job.detected_languages or [])}')
    table=doc.add_table(rows=1, cols=2); table.rows[0].cells[0].text='Speaker label'; table.rows[0].cells[1].text='Имя'
    for a in aliases:
        row=table.add_row().cells; row[0].text=a.speaker_label; row[1].text=a.speaker_name or ''
    doc.add_heading('Стенограмма',1)
    for s in segments: doc.add_paragraph(f'[{format_mmss(s.start_seconds)}–{format_mmss(s.end_seconds)}] {s.speaker_name or s.speaker_label}: {s.text}')
    doc.save(path)
